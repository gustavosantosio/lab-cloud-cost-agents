import json
import asyncio
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
import os
import sys
import hashlib
from pathlib import Path
import pickle

# PDF processing
import PyPDF2
import fitz  # PyMuPDF
from docx import Document

# Vector database and embeddings
import chromadb
from chromadb.config import Settings
import openai
from sentence_transformers import SentenceTransformer

# MCP imports
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import (
    Resource, Tool, TextContent, ImageContent, EmbeddedResource,
    CallToolRequest, CallToolResult, ListResourcesRequest, ListResourcesResult,
    ListToolsRequest, ListToolsResult, ReadResourceRequest, ReadResourceResult
)

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('legal_rag_mcp_server.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class LegalDocument:
    """Documento jurídico"""
    document_id: str
    title: str
    document_type: str  # 'lei', 'decreto', 'instrucao_normativa', 'regulamento'
    source: str
    content: str
    metadata: Dict[str, Any]
    created_at: datetime
    file_path: Optional[str] = None
    page_count: Optional[int] = None


@dataclass
class LegalQuery:
    """Consulta jurídica"""
    query_id: str
    question: str
    context: Dict[str, Any]
    timestamp: datetime


@dataclass
class LegalAnswer:
    """Resposta jurídica"""
    answer_id: str
    query_id: str
    answer: str
    confidence_score: float
    relevant_documents: List[Dict[str, Any]]
    legal_basis: List[str]
    recommendations: List[str]
    risk_assessment: str
    timestamp: datetime


class DocumentProcessor:
    """Processador de documentos"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, int]:
        """Extrai texto de PDF"""
        try:
            text = ""
            page_count = 0
            
            # Tentar com PyMuPDF primeiro (melhor para PDFs complexos)
            try:
                doc = fitz.open(file_path)
                page_count = len(doc)
                
                for page_num in range(page_count):
                    page = doc.load_page(page_num)
                    text += page.get_text()
                
                doc.close()
                
            except Exception as e:
                logger.warning(f"PyMuPDF falhou, tentando PyPDF2: {e}")
                
                # Fallback para PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            
            return text.strip(), page_count
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do PDF {file_path}: {e}")
            return "", 0
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, int]:
        """Extrai texto de DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Contar páginas aproximadamente (250 palavras por página)
            word_count = len(text.split())
            page_count = max(1, word_count // 250)
            
            return text.strip(), page_count
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX {file_path}: {e}")
            return "", 0
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, int]:
        """Extrai texto de TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Contar páginas aproximadamente
            word_count = len(text.split())
            page_count = max(1, word_count // 250)
            
            return text.strip(), page_count
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do TXT {file_path}: {e}")
            return "", 0
    
    def process_document(self, file_path: str, document_type: str = "unknown") -> Optional[LegalDocument]:
        """Processa documento e extrai informações"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                logger.error(f"Arquivo não encontrado: {file_path}")
                return None
            
            extension = file_path.suffix.lower()
            
            if extension not in self.supported_formats:
                logger.error(f"Formato não suportado: {extension}")
                return None
            
            # Extrair texto baseado no formato
            if extension == '.pdf':
                text, page_count = self.extract_text_from_pdf(str(file_path))
            elif extension == '.docx':
                text, page_count = self.extract_text_from_docx(str(file_path))
            elif extension == '.txt':
                text, page_count = self.extract_text_from_txt(str(file_path))
            else:
                return None
            
            if not text:
                logger.error(f"Não foi possível extrair texto de: {file_path}")
                return None
            
            # Gerar ID único baseado no conteúdo
            document_id = hashlib.md5(text.encode()).hexdigest()
            
            # Extrair metadados do texto
            metadata = self._extract_metadata(text, str(file_path))
            
            return LegalDocument(
                document_id=document_id,
                title=metadata.get('title', file_path.stem),
                document_type=document_type,
                source=str(file_path),
                content=text,
                metadata=metadata,
                created_at=datetime.now(),
                file_path=str(file_path),
                page_count=page_count
            )
            
        except Exception as e:
            logger.error(f"Erro ao processar documento {file_path}: {e}")
            return None
    
    def _extract_metadata(self, text: str, file_path: str) -> Dict[str, Any]:
        """Extrai metadados do texto"""
        metadata = {
            'file_path': file_path,
            'file_size': len(text),
            'word_count': len(text.split()),
            'char_count': len(text)
        }
        
        # Tentar identificar tipo de documento baseado no conteúdo
        text_lower = text.lower()
        
        if 'lei n' in text_lower or 'lei nº' in text_lower:
            metadata['document_type'] = 'lei'
        elif 'decreto' in text_lower:
            metadata['document_type'] = 'decreto'
        elif 'instrução normativa' in text_lower or 'instrucao normativa' in text_lower:
            metadata['document_type'] = 'instrucao_normativa'
        elif 'regulamento' in text_lower:
            metadata['document_type'] = 'regulamento'
        elif 'lgpd' in text_lower or 'proteção de dados' in text_lower:
            metadata['document_type'] = 'lgpd'
        elif 'marco civil' in text_lower:
            metadata['document_type'] = 'marco_civil_internet'
        
        # Tentar extrair título
        lines = text.split('\n')[:10]  # Primeiras 10 linhas
        for line in lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 200:
                if any(keyword in line.lower() for keyword in ['lei', 'decreto', 'instrução', 'regulamento']):
                    metadata['title'] = line
                    break
        
        return metadata


class VectorStore:
    """Armazenamento vetorial para documentos"""
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        
        # Configurar ChromaDB
        self.client = chromadb.PersistentClient(
            path=persist_directory,
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True
            )
        )
        
        # Coleção para documentos jurídicos
        self.collection = self.client.get_or_create_collection(
            name="legal_documents",
            metadata={"description": "Documentos jurídicos brasileiros"}
        )
        
        # Modelo de embeddings
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        logger.info(f"Vector store inicializado em: {persist_directory}")
    
    def add_document(self, document: LegalDocument, chunk_size: int = 1000) -> bool:
        """Adiciona documento ao vector store"""
        try:
            # Dividir documento em chunks
            chunks = self._split_text(document.content, chunk_size)
            
            # Gerar embeddings
            embeddings = self.embedding_model.encode(chunks).tolist()
            
            # Preparar metadados para cada chunk
            chunk_ids = []
            chunk_metadatas = []
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{document.document_id}_chunk_{i}"
                chunk_ids.append(chunk_id)
                
                chunk_metadata = {
                    'document_id': document.document_id,
                    'title': document.title,
                    'document_type': document.document_type,
                    'source': document.source,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'created_at': document.created_at.isoformat(),
                    **document.metadata
                }
                chunk_metadatas.append(chunk_metadata)
            
            # Adicionar à coleção
            self.collection.add(
                ids=chunk_ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=chunk_metadatas
            )
            
            logger.info(f"Documento adicionado: {document.title} ({len(chunks)} chunks)")
            return True
            
        except Exception as e:
            logger.error(f"Erro ao adicionar documento ao vector store: {e}")
            return False
    
    def search(self, query: str, n_results: int = 5, 
              document_type: str = None) -> List[Dict[str, Any]]:
        """Busca documentos relevantes"""
        try:
            # Gerar embedding da query
            query_embedding = self.embedding_model.encode([query]).tolist()[0]
            
            # Preparar filtros
            where_filter = {}
            if document_type:
                where_filter['document_type'] = document_type
            
            # Buscar
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                where=where_filter if where_filter else None,
                include=['documents', 'metadatas', 'distances']
            )
            
            # Formatar resultados
            formatted_results = []
            
            if results['documents'] and results['documents'][0]:
                for i in range(len(results['documents'][0])):
                    result = {
                        'document': results['documents'][0][i],
                        'metadata': results['metadatas'][0][i],
                        'similarity_score': 1 - results['distances'][0][i],  # Converter distância para similaridade
                        'relevance_score': (1 - results['distances'][0][i]) * 100
                    }
                    formatted_results.append(result)
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"Erro na busca: {e}")
            return []
    
    def _split_text(self, text: str, chunk_size: int) -> List[str]:
        """Divide texto em chunks"""
        # Dividir por parágrafos primeiro
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # Se o parágrafo sozinho é maior que chunk_size, dividir por sentenças
            if len(paragraph) > chunk_size:
                sentences = paragraph.split('. ')
                for sentence in sentences:
                    if len(current_chunk + sentence) > chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                            current_chunk = sentence
                        else:
                            # Sentença muito longa, dividir por caracteres
                            chunks.extend([sentence[i:i+chunk_size] 
                                         for i in range(0, len(sentence), chunk_size)])
                    else:
                        current_chunk += sentence + ". "
            else:
                if len(current_chunk + paragraph) > chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        current_chunk = paragraph
                    else:
                        chunks.append(paragraph)
                else:
                    current_chunk += paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def get_document_count(self) -> int:
        """Obtém número de documentos"""
        try:
            return self.collection.count()
        except:
            return 0
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """Lista documentos únicos"""
        try:
            # Obter todos os metadados
            results = self.collection.get(include=['metadatas'])
            
            # Agrupar por document_id
            documents = {}
            for metadata in results['metadatas']:
                doc_id = metadata['document_id']
                if doc_id not in documents:
                    documents[doc_id] = {
                        'document_id': doc_id,
                        'title': metadata['title'],
                        'document_type': metadata['document_type'],
                        'source': metadata['source'],
                        'created_at': metadata['created_at'],
                        'total_chunks': metadata.get('total_chunks', 1)
                    }
            
            return list(documents.values())
            
        except Exception as e:
            logger.error(f"Erro ao listar documentos: {e}")
            return []


class LegalRAGSystem:
    """Sistema RAG para documentos jurídicos"""
    
    def __init__(self, documents_directory: str = "./legal_documents", 
                 vector_store_path: str = "./chroma_db"):
        self.documents_directory = Path(documents_directory)
        self.documents_directory.mkdir(exist_ok=True)
        
        self.document_processor = DocumentProcessor()
        self.vector_store = VectorStore(vector_store_path)
        
        # Configurar OpenAI se disponível
        self.openai_client = None
        if os.getenv('OPENAI_API_KEY'):
            openai.api_key = os.getenv('OPENAI_API_KEY')
            self.openai_client = openai
    
    async def load_documents_from_directory(self) -> Dict[str, Any]:
        """Carrega documentos do diretório"""
        try:
            loaded_count = 0
            error_count = 0
            
            for file_path in self.documents_directory.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                    logger.info(f"Processando: {file_path}")
                    
                    document = self.document_processor.process_document(str(file_path))
                    
                    if document:
                        success = self.vector_store.add_document(document)
                        if success:
                            loaded_count += 1
                        else:
                            error_count += 1
                    else:
                        error_count += 1
            
            return {
                'success': True,
                'loaded_documents': loaded_count,
                'errors': error_count,
                'total_documents_in_store': self.vector_store.get_document_count(),
                'timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Erro ao carregar documentos: {e}")
            return {
                'success': False,
                'error': str(e),
                'error_type': 'DOCUMENT_LOADING_ERROR'
            }
    
    async def query_legal_documents(self, question: str, context: Dict[str, Any] = None,
                                   document_type: str = None, max_results: int = 5) -> Dict[str, Any]:
        """Consulta documentos jurídicos"""
        try:
            query_id = hashlib.md5(f"{question}{datetime.now().isoformat()}".encode()).hexdigest()
            
            # Buscar documentos relevantes
            relevant_docs = self.vector_store.search(
                query=question,
                n_results=max_results,
                document_type=document_type
            )
            
            if not relevant_docs:
                return {
                    'success': True,
                    'query_id': query_id,
                    'answer': 'Não foram encontrados documentos relevantes para esta consulta.',
                    'confidence_score': 0.0,
                    'relevant_documents': [],
                    'legal_basis': [],
                    'recommendations': ['Considere reformular a pergunta ou adicionar mais documentos à base.'],
                    'risk_assessment': 'UNKNOWN',
                    'timestamp': datetime.now().isoformat()
                }
            
            # Gerar resposta usando LLM se disponível
            if self.openai_client:
                answer_data = await self._generate_llm_answer(question, relevant_docs, context)
            else:
                answer_data = self._generate_simple_answer(question, relevant_docs)
            
            return {
                'success': True,
                'query_id': query_id,
                'question': question,
                'context': context or {},
                **answer_data,
                'relevant_documents': [
                    {
                        'document_id': doc['metadata']['document_id'],
                        'title': doc['metadata']['title'],
                        'document_type': doc['metadata']['document_type'],
                        'relevance_score': doc['relevance_score'],
                        'excerpt': doc['document'][:500] + "..." if len(doc['document']) > 500 else doc['document']
                    }
                    for doc in relevant_docs
                ],
                'timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Erro na consulta jurídica: {e}")
            return {
                'success': False,
                'error': str(e),
                'error_type': 'LEGAL_QUERY_ERROR'
            }
    
    async def _generate_llm_answer(self, question: str, relevant_docs: List[Dict], 
                                  context: Dict[str, Any] = None) -> Dict[str, Any]:
        """Gera resposta usando LLM"""
        try:
            # Preparar contexto
            context_text = ""
            for i, doc in enumerate(relevant_docs[:3]):  # Top 3 documentos
                context_text += f"\n\nDocumento {i+1} ({doc['metadata']['title']}):\n{doc['document']}"
            
            # Preparar prompt
            system_prompt = """Você é um especialista em direito brasileiro, especialmente em questões relacionadas à tecnologia, proteção de dados e computação em nuvem. 

Analise os documentos fornecidos e responda à pergunta de forma precisa e fundamentada. Sua resposta deve incluir:
1. Uma resposta clara e objetiva
2. Base legal específica (artigos, incisos, parágrafos)
3. Recomendações práticas
4. Avaliação de risco (BAIXO, MÉDIO, ALTO, CRÍTICO)

Seja preciso e cite sempre as fontes."""
            
            user_prompt = f"""Pergunta: {question}

Contexto adicional: {json.dumps(context or {}, indent=2)}

Documentos relevantes:{context_text}

Por favor, forneça uma análise jurídica completa."""
            
            # Chamar OpenAI
            response = await self.openai_client.ChatCompletion.acreate(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1500
            )
            
            answer_text = response.choices[0].message.content
            
            # Extrair componentes da resposta
            legal_basis = self._extract_legal_basis(answer_text)
            recommendations = self._extract_recommendations(answer_text)
            risk_assessment = self._extract_risk_assessment(answer_text)
            
            return {
                'answer': answer_text,
                'confidence_score': min(relevant_docs[0]['similarity_score'] if relevant_docs else 0.5, 0.95),
                'legal_basis': legal_basis,
                'recommendations': recommendations,
                'risk_assessment': risk_assessment,
                'generation_method': 'LLM'
            }
            
        except Exception as e:
            logger.error(f"Erro ao gerar resposta com LLM: {e}")
            return self._generate_simple_answer(question, relevant_docs)
    
    def _generate_simple_answer(self, question: str, relevant_docs: List[Dict]) -> Dict[str, Any]:
        """Gera resposta simples baseada nos documentos"""
        if not relevant_docs:
            return {
                'answer': 'Não foram encontrados documentos relevantes.',
                'confidence_score': 0.0,
                'legal_basis': [],
                'recommendations': [],
                'risk_assessment': 'UNKNOWN',
                'generation_method': 'SIMPLE'
            }
        
        # Combinar trechos mais relevantes
        combined_text = ""
        for doc in relevant_docs[:3]:
            combined_text += f"\n\n{doc['document']}"
        
        answer = f"Com base nos documentos analisados, encontrei as seguintes informações relevantes:\n\n{combined_text[:1000]}..."
        
        return {
            'answer': answer,
            'confidence_score': relevant_docs[0]['similarity_score'],
            'legal_basis': [f"Documento: {doc['metadata']['title']}" for doc in relevant_docs[:2]],
            'recommendations': ['Consulte um advogado especializado para análise detalhada'],
            'risk_assessment': 'MÉDIO',
            'generation_method': 'SIMPLE'
        }
    
    def _extract_legal_basis(self, text: str) -> List[str]:
        """Extrai base legal do texto"""
        import re
        
        patterns = [
            r'[Aa]rt\.?\s*\d+',
            r'[Aa]rtigo\s*\d+',
            r'[Ii]nciso\s*[IVX]+',
            r'[Pp]arágrafo\s*\d+',
            r'[Ll]ei\s*n?º?\s*\d+',
            r'[Dd]ecreto\s*n?º?\s*\d+'
        ]
        
        legal_basis = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            legal_basis.extend(matches)
        
        return list(set(legal_basis))[:5]  # Top 5 únicos
    
    def _extract_recommendations(self, text: str) -> List[str]:
        """Extrai recomendações do texto"""
        lines = text.split('\n')
        recommendations = []
        
        for line in lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in ['recomend', 'sugere', 'deve', 'importante']):
                if len(line) > 20 and len(line) < 200:
                    recommendations.append(line)
        
        return recommendations[:3]  # Top 3
    
    def _extract_risk_assessment(self, text: str) -> str:
        """Extrai avaliação de risco do texto"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['crítico', 'grave', 'severo']):
            return 'CRÍTICO'
        elif any(word in text_lower for word in ['alto', 'elevado', 'significativo']):
            return 'ALTO'
        elif any(word in text_lower for word in ['médio', 'moderado']):
            return 'MÉDIO'
        elif any(word in text_lower for word in ['baixo', 'mínimo', 'reduzido']):
            return 'BAIXO'
        else:
            return 'MÉDIO'


class LegalRAGMCPServer:
    """Servidor MCP para sistema RAG jurídico"""
    
    def __init__(self, documents_directory: str = "./legal_documents"):
        self.server = Server("legal-rag-mcp-server")
        self.rag_system = LegalRAGSystem(documents_directory)
        self._setup_handlers()
    
    def _setup_handlers(self):
        """Configura handlers do servidor MCP"""
        
        @self.server.list_tools()
        async def list_tools() -> List[Tool]:
            """Lista ferramentas disponíveis"""
            return [
                Tool(
                    name="load_legal_documents",
                    description="Carrega documentos jurídicos do diretório",
                    inputSchema={
                        "type": "object",
                        "properties": {},
                        "required": []
                    }
                ),
                Tool(
                    name="query_legal_documents",
                    description="Consulta documentos jurídicos usando RAG",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "question": {"type": "string", "description": "Pergunta jurídica"},
                            "context": {"type": "object", "description": "Contexto adicional"},
                            "document_type": {"type": "string", "description": "Tipo de documento"},
                            "max_results": {"type": "integer", "default": 5}
                        },
                        "required": ["question"]
                    }
                ),
                Tool(
                    name="list_available_documents",
                    description="Lista documentos disponíveis na base",
                    inputSchema={
                        "type": "object",
                        "properties": {},
                        "required": []
                    }
                ),
                Tool(
                    name="get_document_statistics",
                    description="Obtém estatísticas da base de documentos",
                    inputSchema={
                        "type": "object",
                        "properties": {},
                        "required": []
                    }
                )
            ]
        
        @self.server.call_tool()
        async def call_tool(name: str, arguments: Dict[str, Any]) -> CallToolResult:
            """Executa ferramenta"""
            
            try:
                if name == "load_legal_documents":
                    result = await self.rag_system.load_documents_from_directory()
                elif name == "query_legal_documents":
                    result = await self.rag_system.query_legal_documents(**arguments)
                elif name == "list_available_documents":
                    documents = self.rag_system.vector_store.list_documents()
                    result = {
                        'success': True,
                        'documents': documents,
                        'total_count': len(documents),
                        'timestamp': datetime.now().isoformat()
                    }
                elif name == "get_document_statistics":
                    documents = self.rag_system.vector_store.list_documents()
                    
                    # Estatísticas por tipo
                    type_stats = {}
                    for doc in documents:
                        doc_type = doc['document_type']
                        if doc_type not in type_stats:
                            type_stats[doc_type] = 0
                        type_stats[doc_type] += 1
                    
                    result = {
                        'success': True,
                        'statistics': {
                            'total_documents': len(documents),
                            'total_chunks': self.rag_system.vector_store.get_document_count(),
                            'documents_by_type': type_stats,
                            'vector_store_path': self.rag_system.vector_store.persist_directory,
                            'documents_directory': str(self.rag_system.documents_directory)
                        },
                        'timestamp': datetime.now().isoformat()
                    }
                else:
                    result = {"success": False, "error": f"Unknown tool: {name}"}
                
                return CallToolResult(
                    content=[TextContent(
                        type="text",
                        text=json.dumps(result, indent=2, default=str, ensure_ascii=False)
                    )]
                )
                
            except Exception as e:
                logger.error(f"Erro ao executar ferramenta {name}: {e}")
                return CallToolResult(
                    content=[TextContent(
                        type="text",
                        text=json.dumps({
                            "success": False,
                            "error": str(e),
                            "tool": name
                        }, ensure_ascii=False)
                    )]
                )
    
    async def run(self):
        """Executa o servidor MCP"""
        logger.info("Iniciando Legal RAG MCP Server...")
        
        # Executar servidor
        async with stdio_server() as (read_stream, write_stream):
            await self.server.run(read_stream, write_stream)


def main():
    """Função principal"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Legal RAG MCP Server')
    parser.add_argument('--documents-dir', default='./legal_documents', 
                       help='Diretório com documentos jurídicos')
    parser.add_argument('--log-level', default='INFO', help='Nível de log')
    
    args = parser.parse_args()
    
    # Configurar logging
    logging.getLogger().setLevel(getattr(logging, args.log_level.upper()))
    
    # Executar servidor
    server = LegalRAGMCPServer(args.documents_dir)
    
    try:
        asyncio.run(server.run())
    except KeyboardInterrupt:
        logger.info("Servidor interrompido pelo usuário")
    except Exception as e:
        logger.error(f"Erro fatal: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()

